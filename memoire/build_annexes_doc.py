#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Génère Annexes_Hischem_DiagPerf.docx — annexes complètes et autonomes."""
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

NAVY = RGBColor(0x1F, 0x2A, 0x44)
TEAL = RGBColor(0x1D, 0x9E, 0x75)

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(2)

def run(p, t, size=11, color=NAVY, bold=False, italic=False, mono=False):
    r = p.add_run(t); r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name = "Consolas" if mono else "Calibri"
    return r

def H1(t):
    doc.add_page_break()
    p = doc.add_paragraph(); p.space_after = Pt(6)
    run(p, t, 16, NAVY, bold=True)

def H2(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    run(p, t, 12.5, TEAL, bold=True)

def para(t, size=11):
    p = doc.add_paragraph(); run(p, t, size); return p

def code(text):
    for line in text.split("\n"):
        n = len(line) - len(line.lstrip(" "))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
        run(p, (" " * n + line[n:]) if line.strip() else " ", 7.5, RGBColor(0x33, 0x33, 0x33), mono=True)
    doc.add_paragraph()

def img(path, caption):
    w, h = Image.open(path).size; ar = w / h
    iw = 16.0
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if iw / ar > 22:
        p.add_run().add_picture(path, height=Cm(22))
    else:
        p.add_run().add_picture(path, width=Cm(iw))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(c, caption, 9.5, RGBColor(0x55, 0x55, 0x55), italic=True)

def table(header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
    for i, hh in enumerate(header):
        cell = t.rows[0].cells[i]; cell.paragraphs[0].text = ""
        rr = cell.paragraphs[0].add_run(hh); rr.font.bold = True; rr.font.size = Pt(9); rr.font.name = "Calibri"
    for row in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            rc[i].paragraphs[0].text = ""
            rr = rc[i].paragraphs[0].add_run(str(val)); rr.font.size = Pt(8.5); rr.font.name = "Calibri"
    if widths:
        for row in t.rows:
            for i, wdt in enumerate(widths):
                row.cells[i].width = Cm(wdt)
    doc.add_paragraph()

# ── Page de titre ──
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(title, "Annexes", 26, NAVY, bold=True)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(sub, "Mémoire ING2 — « De l'API au RAG » · Cas Diagperf", 13, NAVY)
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(sub2, "Hischem HAMMOUDI — EFREI Paris", 11, RGBColor(0x55, 0x55, 0x55))

# ── Annexe A — Prompts système ──
H1("Annexe A — Prompts système (V1 et V2)")
para("Prompt système de la version V1 (savoir métier baké, sans RAG) :")
code(open("_promptV1.txt", encoding="utf-8").read())
H2("Prompt système V2 (version actuelle, avec contexte RAG injecté)")
code(open("_promptV2.txt", encoding="utf-8").read())

# ── Annexe B — Jeu de test ──
H1("Annexe B — Jeu de test du benchmark (31 questions)")
para("Les 31 questions et leur vérité-terrain. Les cas RC01-RC03 sont transcrits de bugs réels observés en V0/V1.")
qs = json.load(open("../benchmark/questions.json", encoding="utf-8"))["questions"]
rows = []
for q in qs:
    gt = " ".join(str(q.get("ground_truth", "")).split())
    qq = " ".join(str(q.get("question", "")).split())
    rows.append([q["id"], q.get("category", ""), qq, gt])
table(["N°", "Catégorie", "Question", "Vérité-terrain attendue"], rows, widths=[1.4, 2.4, 5.2, 7.0])

# ── Annexe C — Résultats détaillés ──
H1("Annexe C — Résultats détaillés du benchmark")
para("Run du 17 juin 2026 — Claude Haiku 4.5 (générateur), Claude Opus 4.8 (juge). Résultats par condition :")
table(["Condition", "Exactitude", "Hallucinations", "Déterministe OK", "Latence p50"],
      [["A-V1 — baké, sans RAG", "0,77", "9 / 31 (29 %)", "81 %", "2553 ms"],
       ["A-V2 — baké + RAG (livré)", "0,97", "1 / 31", "90 %", "2048 ms"],
       ["B-noRAG — dépouillé, sans RAG", "0,66", "0 / 31", "48 %", "1963 ms"],
       ["B-RAG — dépouillé + RAG", "0,98", "0 / 31", "94 %", "1968 ms"]],
      widths=[6.0, 2.6, 3.0, 3.0, 2.4])
para("Exactitude par famille de questions (A-V1 sans RAG vs B-RAG avec RAG) :")
table(["Famille de questions", "A-V1 (sans RAG)", "B-RAG (avec RAG)"],
      [["Tarifs", "0,75", "1,00"], ["Compatibilité / périmètre", "0,81", "1,00"],
       ["Codes défauts / diagnostic", "1,00", "1,00"], ["Informations / objections", "0,82", "0,95"],
       ["Suivi (questions contextuelles)", "0,33", "1,00"]],
      widths=[8.0, 4.5, 4.5])

# ── Annexe D — Captures d'écran ──
H1("Annexe D — Captures d'écran des conversations")
H2("D.1 — Parcours nominal (V2)")
img("captures/Capture d'écran 2026-06-25 154905.png", "Écran d'accueil : message de bienvenue + bouton « Nos prestations ».")
img("captures/Capture d'écran 2026-06-25 155003.png", "Liste interactive des prestations proposées.")
H2("D.2 — Ruptures de conversation observées en V0/V1")
img("captures/Capture d'écran 2026-06-17 104936.png", "Rupture : « rouler tranquillement ? » → « Je n'ai pas bien saisi votre message ».")
img("captures/Capture d'écran 2026-06-17 104959.png", "Flow à boutons : « risque après ma reprog ? » → « Merci de choisir un des stages ».")
img("captures/Capture d'écran 2026-06-17 105331.png", "Question prise pour une plaque : « reprog pour une essence ? » → « Je n'ai pas reconnu la plaque ».")

# ── Annexe E — Extraits de code ──
H1("Annexe E — Extraits de code clés")
H2("E.1 — Génération d'embedding et expansion de requête (rag.js)")
code('''async function generateEmbedding(text, taskType = "RETRIEVAL_DOCUMENT") {
  const ai = getGoogleAI();
  const result = await ai.models.embedContent({
    model: "gemini-embedding-001",
    contents: text,
    config: { outputDimensionality: 384, taskType },
  });
  return result.embeddings[0].values;
}

function expandQuery(query) {
  const lower = query.toLowerCase().normalize("NFD").replace(/[\\u0300-\\u036f]/g, "");
  const expansions = [];
  for (const [key, synonyms] of Object.entries(QUERY_SYNONYMS)) {
    if (lower.includes(key.normalize("NFD").replace(/[\\u0300-\\u036f]/g, "")))
      expansions.push(synonyms);
  }
  if (expansions.length === 0) return query;
  return `${query} ${expansions.join(" ")}`.slice(0, 512);
}''')
H2("E.2 — Recherche hybride vecteur + plein-texte (PostgreSQL / pgvector)")
code('''CREATE OR REPLACE FUNCTION match_kb_chunks_hybrid(
  query_embedding VECTOR(384), query_text TEXT,
  match_threshold FLOAT DEFAULT 0.2, match_count INT DEFAULT 10,
  keyword_weight FLOAT DEFAULT 0.3
) RETURNS TABLE (...) AS $$
DECLARE ts_query TSQUERY;
BEGIN
  ts_query := plainto_tsquery('french', query_text);
  RETURN QUERY
  SELECT kb.id, kb.content,
    (1 - (kb.embedding <=> query_embedding)) AS similarity,
    COALESCE(ts_rank_cd(kb.fts_content, ts_query), 0) AS keyword_rank,
    ( (1 - keyword_weight) * (1 - (kb.embedding <=> query_embedding))
      + keyword_weight * COALESCE(ts_rank_cd(kb.fts_content, ts_query), 0) * 10
    ) AS combined_score
  FROM kb_chunks kb
  WHERE (1 - (kb.embedding <=> query_embedding)) > match_threshold
     OR kb.fts_content @@ ts_query
  ORDER BY combined_score DESC
  LIMIT match_count;
END; $$ LANGUAGE plpgsql;''')

doc.save("Annexes_Hischem_DiagPerf.docx")
print("OK — Annexes_Hischem_DiagPerf.docx")
