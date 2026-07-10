#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Convertit memoire.md -> memoire.docx (Calibri 12, interligne 1,25, sommaire auto, pagination)."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from PIL import Image
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "memoire.md"
OUT = "Memoire_Hischem_DiagPerf.docx"

doc = Document()

# ---- Styles de base : Calibri 12, interligne 1,25 ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing = 1.25
pf.space_after = Pt(8)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for hname, sz in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
    st = doc.styles[hname]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

# ---- Figures en portrait, intégrées dans le fil du texte (pas de mode paysage) ----
def ensure_landscape():
    return False
def ensure_portrait():
    return False

# ---- Pagination : numéro de page centré en pied ----
def add_page_numbers(document):
    for sec in document.sections:        # numéro sur chaque section (portrait + paysage)
        sec.footer.is_linked_to_previous = False
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

# ---- Champ Sommaire (TOC) que Word met à jour ----
def add_toc(document):
    p = document.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Faites un clic droit puis « Mettre à jour les champs » pour générer le sommaire."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, f2, t, f3):
        run._r.append(e)

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
def add_runs(p, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            p.add_run(tok[2:-2])          # plus de gras dans le texte courant
        elif tok.startswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(11)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])

def add_figure(src, caption):
    try:
        w, h = Image.open(src).size
    except Exception:
        w, h = 4, 3
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (6.3 * h / w) > 8.5:               # figure trop haute -> contrainte par la hauteur
        p.add_run().add_picture(src, height=Inches(8.5))
    else:
        p.add_run().add_picture(src, width=Inches(6.3))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(10)

def add_table(rows):
    ensure_portrait()
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], cells[2:]  # ligne 1 = entête, ligne 2 = séparateur
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        add_runs(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        rc = t.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            rc[i].paragraphs[0].text = ""
            add_runs(rc[i].paragraphs[0], val)
    # éviter qu'un tableau se fractionne en travers d'une page
    keep_together = len(t.rows) <= 7          # petits tableaux : gardés sur une page
    for ri, row in enumerate(t.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))   # une ligne ne se coupe jamais
        for cell in row.cells:
            for cp in cell.paragraphs:
                cp.paragraph_format.keep_together = True
                if keep_together and ri < len(t.rows) - 1:
                    cp.paragraph_format.keep_with_next = True
    doc.add_paragraph()

with open(SRC, encoding="utf-8") as f:
    raw = f.read()

# ---- Page de garde : tout avant "## Remerciements" ----
split = raw.split("## Remerciements", 1)
front, body = split[0], "## Remerciements" + split[1]

for line in front.splitlines():
    s = line.strip()
    if not s or s == "---" or s.startswith(">"):
        continue
    if s.startswith("# "):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s[2:].strip()); r.bold = True; r.font.size = Pt(20)
        p.paragraph_format.space_after = Pt(12)
    else:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        txt = s[2:].strip() if s.startswith("- ") else s
        add_runs(p, txt)
doc.add_page_break()

# ---- Corps ----
lines = body.splitlines()
i = 0
toc_inserted = False
first_h1 = True
para = []

def emit_h1(title):
    global first_h1
    switched = ensure_portrait()
    if not first_h1 and not switched:
        doc.add_page_break()
    first_h1 = False
    doc.add_heading(title, level=1)

def flush():
    global para
    if para:
        ensure_portrait()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        add_runs(p, " ".join(para))
        para = []

def is_break(cs):
    return (cs == "" or cs == "---" or cs.startswith("#") or cs.startswith("|")
            or cs.startswith(">") or cs.startswith("- ") or cs.startswith("![")
            or bool(re.match(r"^\d+\.\s", cs)))

while i < len(lines):
    s = lines[i].strip()
    mimg = re.match(r"!\[(.*?)\]\((.*?)\)", s)
    if mimg:
        flush(); add_figure(mimg.group(2), mimg.group(1)); i += 1; continue
    if s == "[[CODE]]":                       # bloc de code verbatim (monospace)
        flush(); i += 1; code = []
        while i < len(lines) and lines[i].strip() != "[[/CODE]]":
            code.append(lines[i]); i += 1
        i += 1
        for cl in code:
            n = len(cl) - len(cl.lstrip(" "))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run((" " * n + cl[n:]) if cl.strip() else " ")
            r.font.name = "Consolas"; r.font.size = Pt(8)
        doc.add_paragraph()
        continue
    if s.startswith("## "):
        flush(); title = s[3:].strip()
        if title.startswith("1. Introduction") and not toc_inserted:
            emit_h1("Sommaire"); add_toc(doc); toc_inserted = True
        emit_h1(title); i += 1; continue
    if s.startswith("### "):
        flush(); ensure_portrait(); doc.add_heading(s[4:].strip(), level=2); i += 1; continue
    if s.startswith("#### "):
        flush(); ensure_portrait(); doc.add_heading(s[5:].strip(), level=3); i += 1; continue
    if s.startswith("|"):
        flush(); block = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            block.append(lines[i]); i += 1
        add_table(block); continue
    if s.startswith(">"):
        flush(); ensure_portrait(); qbuf = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            qbuf.append(lines[i].strip().lstrip(">").strip()); i += 1
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(24)
        add_runs(p, " ".join(qbuf))
        for r in p.runs:
            r.italic = True
        continue
    if s.startswith("- "):
        flush(); ensure_portrait(); item = [s[2:].strip()]; i += 1
        while i < len(lines) and not is_break(lines[i].strip()):
            item.append(lines[i].strip()); i += 1
        add_runs(doc.add_paragraph(style="List Bullet"), " ".join(item)); continue
    if re.match(r"^\d+\.\s", s):
        flush(); ensure_portrait(); num = [s]; i += 1
        while i < len(lines) and not is_break(lines[i].strip()):
            num.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        add_runs(p, " ".join(num)); continue
    if s == "" or s == "---":
        flush(); i += 1; continue
    para.append(s); i += 1

flush()

add_page_numbers(doc)
# Forcer la mise à jour des champs (sommaire) à l'ouverture, EN RESPECTANT l'ordre du schéma :
# w:updateFields doit précéder w:compat / w:rsids / etc., sinon Word juge le fichier corrompu.
AFTER = {"compat", "rsids", "mathPr", "themeFontLang", "clrSchemeMapping", "shapeDefaults",
         "decimalSymbol", "listSeparator", "hdrShapeDefaults", "footnotePr", "endnotePr",
         "doNotIncludeSubdocsInStats", "uiCompat97To2003", "doNotAutoCompressPictures"}
se = doc.settings.element
upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true")
placed = False
for child in list(se):
    if child.tag.split("}")[-1] in AFTER:
        child.addprevious(upd); placed = True; break
if not placed:
    se.append(upd)
doc.save(OUT)
print("OK ->", OUT)
